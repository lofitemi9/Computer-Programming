# ------------------------------------------------------------
# Name: Temiloluwa Shokunbi
# Student Number: 200632787
# Date: 2025-11-21
# Description:
# The RA Programming & Bulletin‑Board Management System is a Python script that automates tracking of Resident Assistant tasks stored in your “Programming Tracker & Requirements” spreadsheet. It reads each RA’s worksheet and the master sheet, validates inputs, monitors completion of community meetings, bulletin boards, and program requirements, and uses loops and lists to generate reminders for upcoming or overdue tasks. The tool outputs updated reports for managers, ensuring RAs meet programming and community‑building responsibilities

#   RA Programming & Bulletin-Board Management System
#
#   This application reads the "Programming Tracker & Requirements"
#   Excel workbook used by Residence Life. For each RA sheet it:
#     - loads and validates data from Excel
#     - uses loops, lists, string slicing and regex to analyze tasks
#     - checks PERF / evaluation / completion status
#     - builds reminders for incomplete or missing items
#     - writes a Word report for managers (docx)
#     - writes and reads a text summary log file (txt)
# ------------------------------------------------------------

import openpyxl          # Excel
import os                # Paths
import datetime          # Date and time
import time              # time.time()
import pyinputplus as pyip   # Input validation
import docx              # Word docs
import re                # Regular expressions


# ------------------------------------------------------------
# INPUT VALIDATION FOR EXCEL FILE PATH
# ------------------------------------------------------------

def getExcelPath():
    """
    Ask the user for the Excel file path and validate:
      - path exists
      - file has .xlsx extension
    Uses pyinputplus and a while loop (not while True).
    """
    validPath = False
    excelPath = ""

    while not validPath:
        excelPath = pyip.inputFilepath(
            prompt="Enter the full path to the Programming Tracker Excel file:\n"
        )

        if not excelPath.lower().endswith(".xlsx"):
            print("Error: File must be an Excel file (.xlsx). Please try again.")
        elif not os.path.exists(excelPath):
            print("Error: File path does not exist. Please try again.")
        else:
            validPath = True

    return excelPath


# ------------------------------------------------------------
# LOAD EXCEL WORKBOOK (with exception handling)
# ------------------------------------------------------------

def loadExcelFile(filePath):
    """
    Safely load the Excel workbook using openpyxl.
    Returns the workbook object or None if there is an error.
    """
    try:
        workBook = openpyxl.load_workbook(filePath)
        print("Workbook loaded successfully.")
        return workBook
    except FileNotFoundError:
        print("Error: File not found.")
        return None
    except Exception as e:
        print("An unexpected error occurred while loading the workbook:", e)
        return None


# ------------------------------------------------------------
# EXTRACT TASK DATA FROM A SINGLE SHEET
# ------------------------------------------------------------

def extractTasksFromSheet(sheet):
    """
    Read task rows from a worksheet into a list of lists.
    Assumes columns:
      A: Task Name
      B: Program Name
      C: PERF Submitted
      D: PERF Approved
      E: Evaluation Submitted
      F: Evaluation Approved
      G: Completed
      H: Notes
    """
    tasks = []

    for row in sheet.iter_rows(min_row=2, values_only=True):
        # Skip completely empty rows
        allEmpty = True
        for cell in row:
            if cell is not None:
                allEmpty = False
        if allEmpty:
            continue

        task = [
            row[0],  # Task Name
            row[1],  # Program Name
            row[2],  # PERF Submitted
            row[3],  # PERF Approved
            row[4],  # Evaluation Submitted
            row[5],  # Evaluation Approved
            row[6],  # Completed
            row[7]   # Notes
        ]
        tasks.append(task)

    return tasks


# ------------------------------------------------------------
# SIMPLE BOOLEAN CONVERSION
# ------------------------------------------------------------

def toBoolean(value):
    """
    this converys "yes" / "true" style strings into True.
    Any other values (including None) become False.
    """
    if value is None:
        return False

    text = str(value).strip().lower()
    yesValues = ["yes", "y", "true", "t", "1", "done", "complete", "completed"]
    if text in yesValues:
        return True
    else:
        return False


# ------------------------------------------------------------
# CLASSIFY TASK TYPE BASED ON TASK NAME
# ------------------------------------------------------------

def classifyTask(taskName):
    """
    Classify a task into a simple category based on keywords.
    Categories: Program, Community Meeting, Bulletin Board, Other.
    """
    if taskName is None:
        return "Other"

    name = str(taskName).lower()

    if "board" in name:
        return "Bulletin Board"
    elif "community" in name or "cm" in name:
        return "Community Meeting"
    elif "program" in name or "event" in name:
        return "Program"
    else:
        return "Other"


# ------------------------------------------------------------
# ANALYZE TASKS FOR ONE SHEET
#   - String slicing
#   - Regex
# ------------------------------------------------------------

def analyzeTasks(tasks):
    """
    Given a list of task rows, build and return a summary list:
      summary[0] = totalTasks (int)
      summary[1] = completedTasks (int)
      summary[2] = notCompletedTasks (int)
      summary[3] = categoriesSummary (list of [categoryName, total, completed])
      summary[4] = reminders (list of strings)
    """
    totalTasks = 0
    completedTasks = 0
    notCompletedTasks = 0

    # Category counters (lists not needed, just simple ints)
    programTotal = 0
    programCompleted = 0

    communityTotal = 0
    communityCompleted = 0

    boardTotal = 0
    boardCompleted = 0

    otherTotal = 0
    otherCompleted = 0

    reminders = []

    # Regex pattern for a simple date format (YYYY-MM-DD)
    datePattern = r"\d{4}-\d{2}-\d{2}"

    for task in tasks:
        taskName = task[0]
        programName = task[1]
        perfSubmitted = task[2]
        perfApproved = task[3]
        evalSubmitted = task[4]
        evalApproved = task[5]
        completed = task[6]
        notes = task[7]

        totalTasks = totalTasks + 1

        isCompleted = toBoolean(completed)
        isPerfSubmitted = toBoolean(perfSubmitted)
        isPerfApproved = toBoolean(perfApproved)
        isEvalSubmitted = toBoolean(evalSubmitted)
        isEvalApproved = toBoolean(evalApproved)

        category = classifyTask(taskName)

        # Update category counters
        if category == "Program":
            programTotal = programTotal + 1
            if isCompleted:
                programCompleted = programCompleted + 1
        elif category == "Community Meeting":
            communityTotal = communityTotal + 1
            if isCompleted:
                communityCompleted = communityCompleted + 1
        elif category == "Bulletin Board":
            boardTotal = boardTotal + 1
            if isCompleted:
                boardCompleted = boardCompleted + 1
        else:
            otherTotal = otherTotal + 1
            if isCompleted:
                otherCompleted = otherCompleted + 1

        if isCompleted:
            completedTasks = completedTasks + 1
        else:
            notCompletedTasks = notCompletedTasks + 1

        # Choose a label for the program/task
        if programName is not None:
            programLabel = str(programName)
        else:
            programLabel = str(taskName)

        # String slicing example: preview of notes
        notesPreview = ""
        if notes is not None:
            notesText = str(notes)
            notesPreview = notesText[:15]  # first 15 characters

        # Regex example:
        # If the task is completed but notes do NOT contain a date YYYY-MM-DD
        if isCompleted and notes is not None:
            notesTextForRegex = str(notes)
            matchObject = re.search(datePattern, notesTextForRegex)
            if matchObject is None:
                reminders.append(
                    "Consider adding a completion date (YYYY-MM-DD) in the notes for '" +
                    programLabel + "'. Notes preview: '" + notesPreview + "'"
                )

        # PERF-related reminders
        if not isPerfSubmitted:
            reminders.append(
                "Submit PERF for '" + programLabel + "'."
            )
        elif isPerfSubmitted and not isPerfApproved:
            reminders.append(
                "Follow up on PERF approval for '" + programLabel + "'."
            )

        # Completion / evaluation reminders
        if not isCompleted:
            reminders.append(
                "Complete the task '" + programLabel + "'."
            )
        else:
            if not isEvalSubmitted:
                reminders.append(
                    "Submit evaluation for completed task '" + programLabel + "'."
                )
            elif isEvalSubmitted and not isEvalApproved:
                reminders.append(
                    "Follow up on evaluation approval for '" + programLabel + "'."
                )

    # Build categories summary as a list of lists
    categoriesSummary = [
        ["Program", programTotal, programCompleted],
        ["Community Meeting", communityTotal, communityCompleted],
        ["Bulletin Board", boardTotal, boardCompleted],
        ["Other", otherTotal, otherCompleted]
    ]

    summary = [totalTasks, completedTasks, notCompletedTasks, categoriesSummary, reminders]
    return summary


# ------------------------------------------------------------
# SIMPLE DATE/TIME TEXT 
# ------------------------------------------------------------

def getDateTimeParts():
    now = datetime.datetime.now()
    nowText = str(now)
    datePart = nowText[0:10]
    timePart = nowText[11:19]
    return nowText, datePart, timePart


# ------------------------------------------------------------
# BUILD WORD REPORT FOR MANAGERS (docx) - USING LISTS
# ------------------------------------------------------------

def buildManagerReport(summariesBySheet):
    """
    Create a Word document summarizing RA programming status per sheet.
    summariesBySheet is a list of [sheetName, summaryList].
    """
    document = docx.Document()

    # Main heading
    document.add_heading("RA Programming & Bulletin-Board Status Report", level=0)

    # Date / time stamp without strftime
    nowText, datePart, timePart = getDateTimeParts()
    document.add_paragraph("Generated on: " + datePart + " " + timePart[0:5])

    # Loop through each [sheetName, summary] pair
    for item in summariesBySheet:
        sheetName = item[0]
        summary = item[1]

        totalTasks = summary[0]
        completedTasks = summary[1]
        notCompletedTasks = summary[2]
        categoriesSummary = summary[3]
        reminders = summary[4]

        # String slicing on the sheet name, in case it is long
        shortSheetName = sheetName[:25]

        document.add_heading(shortSheetName, level=1)

        statsParagraph = (
            "Total tasks: " + str(totalTasks) + "\n" +
            "Completed: " + str(completedTasks) + "\n" +
            "Not completed: " + str(notCompletedTasks)
        )
        document.add_paragraph(statsParagraph)

        # Breakdown by category
        document.add_paragraph("Breakdown by category:")
        for categoryRow in categoriesSummary:
            categoryName = categoryRow[0]
            categoryTotal = categoryRow[1]
            categoryCompleted = categoryRow[2]

            if categoryTotal > 0:
                line = (
                    " - " + categoryName + ": " +
                    str(categoryCompleted) + " completed out of " +
                    str(categoryTotal) + " tasks"
                )
                document.add_paragraph(line)

        # Reminders
        if len(reminders) > 0:
            document.add_paragraph("Reminders / Follow-ups:")
            for reminder in reminders:
                document.add_paragraph("- " + reminder)
        else:
            document.add_paragraph("No outstanding reminders for this RA.")

    # Build a safe timestamp for filename (no colons)
    nowText, datePart, timePart = getDateTimeParts()
    safeTime = timePart[0:2] + timePart[3:5]
    outputName = "RA_Programming_Report_" + datePart + "_" + safeTime + ".docx"

    try:
        document.save(outputName)
        print("\nManager report created:", outputName)
    except Exception as e:
        print("Error while saving the Word report:", e)


# ------------------------------------------------------------
# WRITE AND READ A TEXT SUMMARY FILE (r/w text files)
# ------------------------------------------------------------

def writeTextSummary(summariesBySheet):
    """
    Write a text file summary (RA_Task_Log.txt) and then
    read back the first few lines to show a preview.
    summariesBySheet is a list of [sheetName, summaryList].
    """
    try:
        nowText, datePart, timePart = getDateTimeParts()

        # Write text file
        logFile = open("RA_Task_Log.txt", "w")
        logFile.write("RA Programming Summary Log\n")
        logFile.write("Generated on: " + datePart + " " + timePart[0:5] + "\n\n")

        for item in summariesBySheet:
            sheetName = item[0]
            summary = item[1]

            totalTasks = summary[0]
            completedTasks = summary[1]
            notCompletedTasks = summary[2]
            reminders = summary[4]

            logFile.write("Sheet: " + sheetName + "\n")
            logFile.write("Total tasks: " + str(totalTasks) + "\n")
            logFile.write("Completed: " + str(completedTasks) + "\n")
            logFile.write("Not completed: " + str(notCompletedTasks) + "\n")
            logFile.write("Number of reminders: " + str(len(reminders)) + "\n")
            logFile.write("-------------------------------------\n\n")

        logFile.close()

        # Read the first few lines back (text file reading)
        previewFile = open("RA_Task_Log.txt", "r")
        print("\nPreview of RA_Task_Log.txt:")
        lineCounter = 0
        for line in previewFile:
            print(line.rstrip("\n"))
            lineCounter = lineCounter + 1
            if lineCounter >= 5:
                break
        previewFile.close()

    except Exception as e:
        print("Error while writing or reading RA_Task_Log.txt:", e)


# ------------------------------------------------------------
# RUN PROGRAM
# ------------------------------------------------------------

def runProgram():
    """
    Run function:
      - Get Excel psth (input validation)
      - load workbook with exception handling
      - loop through sheets and analyze tasks
      - build the word report
      - Write and read text summary file
      - Print total runtime using time module
    """
    startTime = time.time()  

    excelPath = getExcelPath()
    workBook = loadExcelFile(excelPath)

    if workBook is None:
        print("Stopping program because workbook could not be loaded.")
        return

    summariesBySheet = []  # list of [sheetName, summaryList]

    # Loop through each sheet in the workbook
    for sheetName in workBook.sheetnames:
        print("\n-------------------------------------")
        print("Processing sheet:", sheetName)

        sheet = workBook[sheetName]
        tasks = extractTasksFromSheet(sheet)

        # Show first up to 3 tasks for quick verification
        print("First up to 3 tasks from this sheet:")
        index = 0
        for task in tasks:
            print(task)
            index = index + 1
            if index >= 3:
                break

        summary = analyzeTasks(tasks)
        summariesBySheet.append([sheetName, summary])

    # Build Word report
    buildManagerReport(summariesBySheet)

    # Write and read text summary file
    writeTextSummary(summariesBySheet)

    endTime = time.time()
    elapsedSeconds = endTime - startTime
    print("\nThis report took " + str(elapsedSeconds) + " seconds to generate.")

runProgram()

